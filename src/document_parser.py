"""
Document Parser Module
Handles parsing of multiple document formats: PDF, TXT, MD, DOCX
"""

import os
import re
from pathlib import Path
from typing import Optional

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

import markdown
from bs4 import BeautifulSoup


class DocumentParser:
    """Parse various document formats and extract clean text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx']
    
    def parse(self, file_path: str) -> str:
        """
        Parse a document and return its text content
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted and cleaned text
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {self.supported_formats}")
        
        # Parse based on file type
        if suffix == '.pdf':
            text = self._parse_pdf(file_path)
        elif suffix == '.txt':
            text = self._parse_txt(file_path)
        elif suffix == '.md':
            text = self._parse_markdown(file_path)
        elif suffix == '.docx':
            text = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")
        
        # Clean the extracted text
        return self._clean_text(text)
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file"""
        if PdfReader is None:
            raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
        
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            return text
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown file and convert to plain text"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html = markdown.markdown(md_content)
        
        # Extract text from HTML
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n\n')
        
        return text
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        if Document is None:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and special characters
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove any remaining weird characters (but keep basic punctuation)
        text = re.sub(r'[^\w\s\.,;:!?\-\'"()\[\]{}/@#$%&*+=<>]', '', text)
        
        return text.strip()
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get basic information about the file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        path = Path(file_path)
        
        return {
            'filename': path.name,
            'format': path.suffix,
            'size_bytes': path.stat().st_size,
            'size_kb': round(path.stat().st_size / 1024, 2),
        }


# Example usage
if __name__ == "__main__":
    parser = DocumentParser()
    
    # Test with a sample file
    test_file = "test.txt"
    
    if os.path.exists(test_file):
        text = parser.parse(test_file)
        print(f"Extracted {len(text)} characters")
        print(f"\nFirst 500 characters:\n{text[:500]}")
    else:
        print(f"Test file '{test_file}' not found")

